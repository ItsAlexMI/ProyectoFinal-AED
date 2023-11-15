from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session, send_file
from tempfile import mkdtemp
from flask_session import Session
from cs50 import SQL
from werkzeug.security import *
from helpers import *
import time
from flask_excel import make_response_from_array
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.lib.pagesizes import letter
from docx import Document
from pyexcel_xlsx import save_data
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config["SESSION_FILE_DIR"] = mkdtemp()
app.config['DEBUG'] = True
app.config["SESSION_TYPE"] = "filesystem"
Session(app)
db = SQL("sqlite:///database.db")
app.config['PRODUCT_IMAGES_FOLDER'] = 'static/product_images/'
app.config['USER_IMAGES_FOLDER'] = 'static/user_images/'

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def capturar_username():
    user_id = session.get("user_id")

    user = db.execute("SELECT NombreUsuario FROM Usuarios WHERE id = ?", user_id)

    username = user[0]["NombreUsuario"] if user else None

    return username

def capturar_correo():
    user_id = session.get("user_id")

    user = db.execute("SELECT CorreoElectronico FROM Usuarios WHERE id = ?", user_id)

    correo = user[0]["CorreoElectronico"] if user else None

    return correo

def capturar_foto():
    user_id = session.get("user_id")

    user = db.execute("SELECT UrlImagenPerfil FROM Usuarios WHERE id = ?", user_id)

    url_imagen_perfil = user[0]["UrlImagenPerfil"] if user else None


    return url_imagen_perfil
def actualizar_producto(producto_id, CategoriaID, NombreProducto, Cantidad, PrecioOriginal, PrecioVenta, FechaActualizacion, Imagen):
    db.execute("UPDATE Inventario SET CategoriaID = :CategoriaID, NombreProducto = :NombreProducto, Cantidad = :Cantidad, PrecioOriginal = :PrecioOriginal, PrecioVenta = :PrecioVenta, FechaActualizacion = :FechaActualizacion, Imagen = :Imagen WHERE ProductoID = :producto_id",
               CategoriaID=CategoriaID, NombreProducto=NombreProducto, Cantidad=Cantidad, PrecioOriginal=PrecioOriginal, PrecioVenta=PrecioVenta, FechaActualizacion=FechaActualizacion, producto_id=producto_id, Imagen=Imagen)

@app.route('/download_table/<format>', methods=['GET'])
def download_table(format):
    if format == 'pdf':
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        p.setFont("Helvetica", 12)
        p.drawString(100, 750, "Tabla de Ventas")

        user_id = session.get('user_id')
        ventas = db.execute("SELECT FechaVenta, NombreProductoVendido, TotalVentas, CantidadVendida FROM Ventas WHERE UsuarioID = ?", user_id)

        y = 700
        for venta in ventas:
            y -= 20
            p.drawString(100, y, f"Producto: {venta['NombreProductoVendido']}")
            y -= 20
            p.drawString(100, y, f"Fecha de venta: {venta['FechaVenta']}")
            y -= 20
            p.drawString(100, y, f"Cantidad vendida: {venta['CantidadVendida']}")
            y -= 20
            p.drawString(100, y, f"Ganancias totales de la venta: {venta['TotalVentas']}")

        p.showPage()
        p.save()
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='tabla.pdf', mimetype='application/pdf') 
    elif format == 'excel':
        user_id = session.get('user_id')
        ventas = db.execute("SELECT FechaVenta, NombreProductoVendido, TotalVentas, CantidadVendida FROM Ventas WHERE UsuarioID = ?", user_id)
        array = [['Producto', 'Fecha de venta', 'Cantidad vendida', 'Ganancias totales de la venta']]
        for venta in ventas:
            array.append([venta['NombreProductoVendido'], venta['FechaVenta'], venta['CantidadVendida'], venta['TotalVentas']])

        excel_file_path = 'tabla.xlsx'
        save_data(excel_file_path, {"Hoja 1": array})

        return send_file(excel_file_path, as_attachment=True, download_name='tabla.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        return response
    elif format == 'word':
        user_id = session.get('user_id')
        ventas = db.execute("SELECT FechaVenta, NombreProductoVendido, TotalVentas, CantidadVendida FROM Ventas WHERE UsuarioID = ?", user_id)
        
        doc = Document()
        doc.add_heading('Tabla de Ventas', 0)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Producto'
        heading_cells[1].text = 'Fecha de venta'
        heading_cells[2].text = 'Cantidad vendida'
        heading_cells[3].text = 'Ganancias totales de la venta'
        
        for venta in ventas:
            row_cells = table.add_row().cells
            row_cells[0].text = venta['NombreProductoVendido']
            row_cells[1].text = venta['FechaVenta']
            row_cells[2].text = str(venta['CantidadVendida'])
            row_cells[3].text = str(venta['TotalVentas'])
        
        word_file_path = 'tabla.docx'
        doc.save(word_file_path)
        
        return send_file(word_file_path, as_attachment=True, download_name='tabla.docx', mimetype='application/msword')

    return "Formato no válido"

@app.route('/Fechas', methods=['GET'])
def Fechas():
    user_id = session.get('user_id')
 
    
    resultados = db.execute("SELECT FechaVenta, SUM(CantidadVendida) AS Fechas FROM Ventas WHERE UsuarioID = ? GROUP BY FechaVenta ORDER BY Fechas DESC LIMIT 5", user_id)

    Fechas = [{"FechaVenta": row["FechaVenta"], "CantidadVendida": row["Fechas"]} for row in resultados]

    return jsonify(Fechas)

@app.route('/ProductoMasVendido', methods=['GET'])
def productos_mas_vendidos():
    user_id = session.get('user_id')
    
    resultados = db.execute("SELECT NombreProductoVendido, COUNT(CantidadVendida) AS ProductoMasVendido FROM Ventas WHERE UsuarioID = ? GROUP BY NombreProductoVendido ORDER BY ProductoMasVendido DESC LIMIT 5", user_id)

    productos_mas_vendidos = [{"NombreProductoVendido": row["NombreProductoVendido"], "CantidadVendida": row["ProductoMasVendido"]} for row in resultados]

    return jsonify(productos_mas_vendidos)

def obtener_producto(producto_id):
    producto = db.execute("SELECT * FROM Inventario WHERE ProductoID = :producto_id", producto_id=producto_id)
    
    if producto:
        return producto[0] 
    else:
        return None 

@app.route('/Stock', methods=['GET'])
def obtenerStock():

    user_id = session.get('user_id')

    resultados = db.execute("SELECT NombreProducto, Cantidad FROM Inventario WHERE UsuarioID = ? ORDER BY Cantidad DESC LIMIT 5", user_id)
    
    productos_mas_vendidos = [{"NombreProducto": row["NombreProducto"], "Cantidad": row["Cantidad"]} for row in resultados]
    
    return jsonify(productos_mas_vendidos)

@app.route('/Ganancias', methods=['GET'])
def obtenerGanancia():

    user_id = session.get('user_id')

    resultados = db.execute("SELECT NombreProductoVendido, SUM(TotalVentas) AS GananciasTotales FROM Ventas WHERE UsuarioID = ? GROUP BY NombreProductoVendido ORDER BY GananciasTotales DESC LIMIT 5", user_id)

    productos_mas_ganancias = [{"NombreProductoVendido": row["NombreProductoVendido"], "TotalVentas": row["GananciasTotales"]} for row in resultados]

    return jsonify(productos_mas_ganancias)



@app.route('/loading')
@login_required
def loading():
    time.sleep(6)  
    return render_template('load.html')
@app.route("/")
@login_required
def index():
    url_imagen_perfil = capturar_foto()
    productos_mas_vendidos = obtenerStock()
    username = capturar_username()
    correo = capturar_correo()

    


    

    UsuarioID = session['user_id']
    productos_propios = db.execute("""
    SELECT 
        I.ProductoID,
        I.NombreProducto, 
        I.Cantidad, 
        I.FechaActualizacion, 
        I.PrecioOriginal, 
        I.PrecioVenta, 
        I.Imagen,
        C.NombreCategoria, 
        C.Descripcion 
    FROM Inventario AS I
    LEFT JOIN Categorias AS C ON I.CategoriaID = C.ID
    WHERE I.UsuarioID = ?
    """, UsuarioID)

    gananciasTotales = db.execute("SELECT SUM(TotalVentas) AS GananciaTotal FROM Ventas WHERE UsuarioID = ?", UsuarioID)
    cantidadProductos = db.execute("SELECT COUNT(*) AS Cantidad FROM Inventario WHERE UsuarioID = ?", UsuarioID)
    ventasTotales =  db.execute("SELECT count(*) AS CantidadVendida FROM Ventas WHERE UsuarioID = ?", UsuarioID)


    if productos_mas_vendidos is not None:
        return render_template('index.html', correo=correo ,productos_mas_vendidos=productos_mas_vendidos, url_imagen_perfil=url_imagen_perfil, productos_propios=productos_propios, gananciasTotales=gananciasTotales, cantidadProductos=cantidadProductos, ventasTotales=ventasTotales, username = username)
        return render_template('index.html', correo=correo , url_imagen_perfil=url_imagen_perfil, productos_propios=productos_propios, gananciasTotales=gananciasTotales , cantidadProductos=cantidadProductos, ventasTotales=ventasTotales, username=username)
    

@app.route("/cambiar-foto", methods=["GET", "POST"])
@login_required
def cambiar_foto():
    username = capturar_username()
    correo = capturar_correo()
    url_imagen_perfil = capturar_foto()

    if request.method == "POST":
        imagen_perfil = request.files['imagen_perfil']

        if imagen_perfil.filename != '':
            if imagen_perfil and allowed_file(imagen_perfil.filename):
                filename = secure_filename(imagen_perfil.filename)
                if not os.path.exists(app.config['USER_IMAGES_FOLDER']):
                    os.makedirs(app.config['USER_IMAGES_FOLDER'])
                imagen_perfil.save(os.path.join(app.config['USER_IMAGES_FOLDER'], filename))
                url_imagen_perfil = os.path.join(app.config['USER_IMAGES_FOLDER'], filename)

                db.execute("UPDATE Usuarios SET UrlImagenPerfil = ? WHERE id = ?", url_imagen_perfil, session['user_id'])
                return redirect(url_for("index"))

    return render_template('cambiarFoto.html', url_imagen_perfil=url_imagen_perfil, username=username, correo=correo)
@app.route("/cambiar-datos", methods=["GET", "POST"])
@login_required
def cambiar_datos():
    url_imagen_perfil = capturar_foto()
    username = capturar_username()
    correo = capturar_correo()

    if request.method == "POST":
        new_username = request.form.get("username")
        new_correo = request.form.get("email")

        if new_username:
            username = new_username

        if new_correo:
            correo = new_correo

        if new_username or new_correo:
            db.execute("UPDATE Usuarios SET NombreUsuario = ?, CorreoElectronico = ? WHERE id = ?", username, correo, session['user_id'])
            return redirect(url_for("index"))

    return render_template('cambiarDatos.html', username=username, correo=correo, url_imagen_perfil=url_imagen_perfil)



#mezclar el index con graficas
@app.route("/dash")
@login_required
def dash():
    url_imagen_perfil = capturar_foto()
    username = capturar_username()
    correo = capturar_correo()
    # productos_mas_vendidos = obtenerStock()
    UsuarioID = session['user_id']
    productos_propios = db.execute("""
    SELECT 
        I.ProductoID,
        I.NombreProducto, 
        I.Cantidad, 
        I.FechaActualizacion, 
        I.PrecioOriginal, 
        I.PrecioVenta, 
        C.NombreCategoria, 
        C.Descripcion 
    FROM Inventario AS I
    LEFT JOIN Categorias AS C ON I.CategoriaID = C.ID
    WHERE I.UsuarioID = ?
    """, UsuarioID)

    stock = db.execute("SELECT Cantidad AS Stock FROM Inventario WHERE UsuarioID = ? ORDER BY Cantidad DESC LIMIT 1", UsuarioID)
    producto_mas_ganancias = db.execute("SELECT NombreProductoVendido AS ProductoTop FROM Ventas WHERE UsuarioID = ? ORDER BY TotalVentas DESC LIMIT 1", UsuarioID)
    productos_mas_vendidos = db.execute("SELECT NombreProductoVendido AS Producto FROM Ventas WHERE UsuarioID = ? ORDER BY CantidadVendida DESC LIMIT 1", UsuarioID)    
    producto_menos_vendido = db.execute("SELECT NombreProductoVendido AS ProductoMV FROM Ventas WHERE UsuarioID = ? ORDER BY CantidadVendida ASC LIMIT 1", UsuarioID)
    producto_menos_ganancia = db.execute("SELECT NombreProductoVendido AS ProductoMG FROM Ventas WHERE UsuarioID = ? ORDER BY TotalVentas ASC LIMIT 1", UsuarioID)
    gananciasTotales = db.execute("SELECT SUM(TotalVentas) AS GananciaTotal FROM Ventas WHERE UsuarioID = ?", UsuarioID)
    cantidadProductos = db.execute("SELECT COUNT(*) AS Cantidad FROM Inventario WHERE UsuarioID = ?", UsuarioID)
    ventasTotales =  db.execute("SELECT count(*) AS CantidadVendida FROM Ventas WHERE UsuarioID = ?", UsuarioID)
    ventas = db.execute("SELECT FechaVenta, NombreProductoVendido, TotalVentas, CantidadVendida FROM Ventas WHERE UsuarioID = ?", UsuarioID)



    if productos_mas_vendidos is not None:
        return render_template('dash.html', correo=correo ,productos_mas_vendidos=productos_mas_vendidos, url_imagen_perfil=url_imagen_perfil, productos_propios=productos_propios, gananciasTotales=gananciasTotales, cantidadProductos=cantidadProductos, ventasTotales=ventasTotales, username = username, ventas = ventas, producto_mas_ganancias=producto_mas_ganancias, stock = stock, producto_menos_vendido=producto_menos_vendido, producto_menos_ganancia=producto_menos_ganancia)
        return render_template('dash.html', correo=correo , url_imagen_perfil=url_imagen_perfil, productos_propios=productos_propios, gananciasTotales=gananciasTotales , cantidadProductos=cantidadProductos, ventasTotales=ventasTotales, username=username, ventas = ventas , producto_mas_ganancia=producto_mas_ganancia, stock = stock, producto_menos_vendido=producto_menos_vendido, producto_menos_ganancia=producto_menos_ganancia)
 

@app.route("/graficas")
@login_required
def graficas():
    url_imagen_perfil = capturar_foto()
    ganancias = obtenerGanancia()
    username = capturar_username()
    correo = capturar_correo()
    UsuarioID = session['user_id']
    productos_propios = db.execute("""
    SELECT 
        I.ProductoID,
        I.NombreProducto, 
        I.Cantidad, 
        I.FechaActualizacion, 
        I.PrecioOriginal, 
        I.PrecioVenta, 
        C.NombreCategoria, 
        C.Descripcion 
    FROM Inventario AS I
    LEFT JOIN Categorias AS C ON I.CategoriaID = C.ID
    WHERE I.UsuarioID = ?
    """, UsuarioID)

    ventas = db.execute("SELECT FechaVenta, NombreProductoVendido, TotalVentas, CantidadVendida FROM Ventas WHERE UsuarioID = ?", UsuarioID)

    if ganancias is not None:
        return render_template('graficas.html', username = username, correo = correo ,ganancias=ganancias, url_imagen_perfil=url_imagen_perfil, productos_propios=productos_propios, ventas=ventas)
    else:
        return render_template('graficas.html', username = username, correo = correo ,url_imagen_perfil=url_imagen_perfil, productos_propios=productos_propios, ventas=ventas)

@app.route("/editar-producto/<int:producto_id>", methods=["GET", "POST"])
@login_required
def editar_producto(producto_id):
    url_imagen_perfil = capturar_foto()
    categorias = db.execute("SELECT ID, NombreCategoria FROM Categorias")
    username = capturar_username()
    correo = capturar_correo()
    
    producto = obtener_producto(producto_id)
    if request.method == "POST":
        CategoriaID = request.form.get("CategoriaID")
        NombreProducto = request.form.get("NombreProducto")
        Cantidad = request.form.get("Cantidad")
        PrecioOriginal = request.form.get("PrecioOriginal")
        PrecioVenta = request.form.get("PrecioVenta")
        FechaActualizacion = request.form.get("FechaActualizacion")
        categoria = request.form.get("categoria")
        descripcion = request.form.get("descripcion")

        imagen = request.files['imagen_producto']

        if imagen.filename != '':
            if imagen and allowed_file(imagen.filename):
                filename = secure_filename(imagen.filename)
                if not os.path.exists(app.config['PRODUCT_IMAGES_FOLDER']):
                    os.makedirs(app.config['PRODUCT_IMAGES_FOLDER'])
                imagen.save(os.path.join(app.config['PRODUCT_IMAGES_FOLDER'], filename))
                url_imagen_producto = os.path.join(app.config['PRODUCT_IMAGES_FOLDER'], filename)

                actualizar_producto(producto_id, CategoriaID, NombreProducto, Cantidad, PrecioOriginal, PrecioVenta, FechaActualizacion, url_imagen_producto)
            else:
                flash('Formato de imagen no válido. Sube imágenes en formato PNG, JPG, JPEG o GIF.')

        else:
            url_imagen_producto = producto['Imagen']
            actualizar_producto(producto_id, CategoriaID, NombreProducto, Cantidad, PrecioOriginal, PrecioVenta, FechaActualizacion)

        return redirect(url_for("index"))
    
    return render_template("editarProducto.html", username=username, correo=correo , categorias=categorias, producto=producto, url_imagen_perfil=url_imagen_perfil)
    
@app.route("/buscar-productos", methods=["GET"])
@login_required
def buscar_productos():
    
    url_imagen_perfil = capturar_foto()
    username = capturar_username()
    correo = capturar_correo()
    query = request.args.get("query")

    if query:
        productos_encontrados = db.execute(
            "SELECT * FROM Inventario WHERE NombreProducto LIKE ?", f"%{query}%"
        )
    else:
        productos_encontrados = []

    return render_template("resultados_busqueda.html", productos=productos_encontrados, query=query, url_imagen_perfil=url_imagen_perfil, username=username, correo=correo)

@app.route("/eliminar-producto/<int:producto_id>")
@login_required
def eliminar_producto(producto_id):
    db.execute("DELETE FROM Inventario WHERE ProductoID = ?", producto_id)

    return redirect(url_for("index"))



@app.route("/login", methods=["GET", "POST"])

def login():

    session.clear()

    if request.method == "POST":
        email = request.form.get("email-login")
        password = request.form.get("password-login")

        users = db.execute("SELECT * FROM Usuarios WHERE CorreoElectronico = ?", email)
        if not users:
            flash("El usuario no existe")
            return redirect(url_for("login"))

        user = users[0]

        if check_password_hash(user["Contraseña"], password):
            print("Contraseña correcta")
            session["user_id"] = user["id"]
            return redirect(url_for("loading"))
        else:
            print("Contraseña incorrecta")
            flash("Contraseña incorrecta")
            return redirect(url_for("login"))


            

    return render_template("login.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = request.form.get("nombre")
        email = request.form.get("email-register")
        password = request.form.get("password-register")

        hashpass = generate_password_hash(password)

        result = db.execute("SELECT * FROM Usuarios WHERE CorreoElectronico = ?", email)
        user = result[0] if result else None

        if user is not None:
            flash("El usuario ya existe")
            return redirect(url_for("register"))

        db.execute("INSERT INTO Usuarios (NombreUsuario, CorreoElectronico, Contraseña) VALUES (?, ?, ?)", name, email, hashpass)

        return jsonify({"success": True, "message": "¡Usuario registrado con éxito!"})

    return redirect("/login")

@app.route("/logout")
def logout():
    session.clear()
    return render_template("loadlogout.html")


@app.route("/agregar-producto", methods=["GET", "POST"])
@login_required
def agregar_producto():
    url_imagen_perfil = capturar_foto()
    username = capturar_username()
    correo = capturar_correo()
    categorias = db.execute("SELECT ID, NombreCategoria FROM Categorias")
    
    if request.method == "POST":
        if 'user_id' not in session:
            return redirect(url_for('login'))

        UsuarioID = session['user_id']
        CategoriaID = request.form.get("CategoriaID")
        NombreProducto = request.form.get("NombreProducto")
        Cantidad = request.form.get("Cantidad")
        PrecioOriginal = request.form.get("PrecioOriginal")
        PrecioVenta = request.form.get("PrecioVenta")
        FechaActualizacion = request.form.get("FechaActualizacion")
        
        imagen = request.files.get('imagen_producto')  

        if imagen and allowed_file(imagen.filename):
            filename = secure_filename(imagen.filename)
            if not os.path.exists(app.config['PRODUCT_IMAGES_FOLDER']):
                os.makedirs(app.config['PRODUCT_IMAGES_FOLDER'])
            imagen.save(os.path.join(app.config['PRODUCT_IMAGES_FOLDER'], filename))
            url_imagen_producto = os.path.join(app.config['PRODUCT_IMAGES_FOLDER'], filename)


            db.execute("INSERT INTO Inventario (UsuarioID, CategoriaID, NombreProducto, Cantidad, PrecioOriginal, PrecioVenta, FechaActualizacion, Imagen) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                       UsuarioID, CategoriaID, NombreProducto, Cantidad, PrecioOriginal, PrecioVenta, FechaActualizacion, url_imagen_producto)

            return redirect(url_for("index"))
        elif imagen:
            flash('Formato de imagen no válido. Sube imágenes en formato PNG, JPG, JPEG o GIF.')

    return render_template("agregarProducto.html", username=username, correo=correo, categorias=categorias, url_imagen_perfil=url_imagen_perfil)
@app.route("/vender-producto", methods=["GET", "POST"])
def vender_producto():
    url_imagen_perfil = capturar_foto()
    username = capturar_username()
    correo = capturar_correo()

    if request.method == "POST":
        if 'user_id' not in session:
            return redirect(url_for('login'))

        UsuarioID = session['user_id']
        ProductoID = request.form.get("ProductoID")
        CantidadVendida = int(request.form.get("CantidadVendida"))

        producto = db.execute("SELECT NombreProducto, PrecioOriginal, PrecioVenta, Cantidad FROM Inventario WHERE ProductoID = ? AND UsuarioID = ?", ProductoID, UsuarioID)

        if producto:
            nombre_producto = producto[0]["NombreProducto"]
            precio_original = producto[0]["PrecioOriginal"]
            precio_venta = producto[0]["PrecioVenta"]
            cantidad_actual = producto[0]["Cantidad"]

            if CantidadVendida <= cantidad_actual:
                ganancia_total = (precio_venta - precio_original) * CantidadVendida
                nueva_cantidad = cantidad_actual - CantidadVendida
                db.execute("UPDATE Inventario SET Cantidad = ? WHERE ProductoID = ? AND UsuarioID = ?", nueva_cantidad, ProductoID, UsuarioID)

                db.execute("INSERT INTO Ventas (FechaVenta, UsuarioID, NombreProductoVendido, TotalVentas, CantidadVendida) VALUES (CURRENT_DATE, ?, ?, ?, ?)",
                           UsuarioID, nombre_producto, ganancia_total, CantidadVendida)

                return jsonify({"success": True, "message": "¡Producto vendido con éxito!"})
            else:
                return jsonify({"error": True, "message": "No hay suficiente cantidad en el inventario para vender."})
        else:
            return jsonify({"error": True, "message": "No se encontró el producto."})

    UsuarioID = session['user_id']
    inventario = db.execute("SELECT ProductoID, NombreProducto FROM Inventario WHERE UsuarioID = ?", UsuarioID)

    return render_template("venderProducto.html", username=username, correo=correo, inventario=inventario, url_imagen_perfil=url_imagen_perfil)